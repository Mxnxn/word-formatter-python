from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
import glob
import os

# open folder of docs files
os.chdir("/Himani Docs")

# iterate over multiple docx files
for file in glob.glob("*.docx"):
    # Reading of docx file
    doc = Document(file)

    # iterate each paragraph to change its fonts and size
    for para in doc.paragraphs:
        for run in para.runs:
            # Any word font eg. Times New Roman, Arial, Courier New
            run.font.name = "Calibri"
            # Size of each font
            run.font.size = Pt(14)

    # iterate over table to row to cells for changing fonts and providing borders
    for table in doc.tables:
        for i in range(0, len(table.rows)):
            for row in table.rows:
                for cell in row.cells:
                    #! TO FILL BACKGROUND
                    # shading_elm = parse_xml(
                    #     r'<w:shd {} w:val="horzStripe"/>'.format(nsdecls('w')))
                    # <w:top/> will provide top border of cell plus w:sz for bordersize, w:space for spacing, w:color for border color
                    # same will goes for w:start,w:end,W:bottom
                    shading_elm = parse_xml(
                        r'<w:top {} w:val="single" w:sz="5" w:space="0" w:color="000000" />'.format(nsdecls('w')))
                    shading_elm2 = parse_xml(
                        r'<w:start {} w:val="single" w:sz="5" w:space="0" w:color="000000" />'.format(nsdecls('w')))
                    shading_elm3 = parse_xml(
                        r'<w:end {} w:val="single" w:sz="5" w:space="0" w:color="000000" />'.format(nsdecls('w')))
                    shading_elm4 = parse_xml(
                        r'<w:bottom {} w:val="single" w:sz="5" w:space="0" w:color="000000" />'.format(nsdecls('w')))
                    # appending XMLs to cell xml
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    cell._tc.get_or_add_tcPr().append(shading_elm2)
                    cell._tc.get_or_add_tcPr().append(shading_elm3)
                    cell._tc.get_or_add_tcPr().append(shading_elm4)

                    # to change its font and style
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(14)

    # basically overwriting file, change Here to create new with formating
    doc.save(file)
